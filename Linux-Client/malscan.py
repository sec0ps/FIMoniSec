# =============================================================================
# FIMonsec Tool - File Integrity Monitoring Security Solution
# =============================================================================
#
# Author: Keith Pachulski
# Company: Red Cell Security, LLC
# Email: keith@redcellsecurity.org
# Website: www.redcellsecurity.org
#
# Copyright (c) 2025 Keith Pachulski. All rights reserved.
#
# License: This software is licensed under the MIT License.
#          You are free to use, modify, and distribute this software
#          in accordance with the terms of the license.
#
# Purpose: This script is part of the FIMonsec Tool, which provides enterprise-grade
#          system integrity monitoring with real-time alerting capabilities. It monitors
#          critical system and application files for unauthorized modifications,
#          supports baseline comparisons, and integrates with SIEM solutions.
#
# DISCLAIMER: This software is provided "as-is," without warranty of any kind,
#             express or implied, including but not limited to the warranties
#             of merchantability, fitness for a particular purpose, and non-infringement.
#             In no event shall the authors or copyright holders be liable for any claim,
#             damages, or other liability, whether in an action of contract, tort, or otherwise,
#             arising from, out of, or in connection with the software or the use or other dealings
#             in the software.
#
# =============================================================================
import os
import re
import math
import json
import csv
import hashlib
import logging
import argparse
import chardet
import zipfile
import mimetypes
import tempfile
import datetime
from malscan_yara import yara_scan_file
import xml.etree.ElementTree as ET
from typing import List, Tuple, Union, Dict, Any
from concurrent.futures import ThreadPoolExecutor, as_completed
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from pdfminer.high_level import extract_text
from email import policy
from email.parser import BytesParser
from email.message import EmailMessage
from docx import Document
import openpyxl
import requests
import warnings

# Filter out XML parsed as HTML warning
warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('malware_scanner.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Updated ScannerConfig class
class ScannerConfig:
    """Centralized configuration for the malware scanner"""
    ENTROPY_THRESHOLD = 7.5
    VIRUS_TOTAL_API_KEY = os.environ.get('VIRUSTOTAL_API_KEY', '')
    MAX_FILE_SIZE_MB = 50  # Maximum file size to scan
    ENABLED_CHECKS = {
        'entropy': True,
        'indicators': True,
        'virustotal': bool(VIRUS_TOTAL_API_KEY),
        'yara': True
    }
    # Default YARA rules directory - now with path resolution
    YARA_RULES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'yara_rules')
    
    @classmethod
    def create_default_directories(cls):
        """Create default directories if they don't exist"""
        # Create YARA rules directory if it doesn't exist
        if not os.path.exists(cls.YARA_RULES_DIR):
            try:
                os.makedirs(cls.YARA_RULES_DIR)
                logger.info(f"Created YARA rules directory at {cls.YARA_RULES_DIR}")
                
                # Create a sample YARA rule for demonstration
                sample_rule = """
rule SuspiciousFile {
    meta:
        description = "Detects suspicious file characteristics"
        author = "Scanner Framework"
        severity = "medium"
        date = "2025-04-02"
    
    strings:
        $s1 = "cmd.exe" nocase
        $s2 = "powershell" nocase
        $s3 = "wget" nocase
        $s4 = "curl" nocase
        $exec1 = "Execute" nocase
        $exec2 = "eval(" nocase
        $exec3 = "exec(" nocase

    condition:
        2 of ($s*) or any of ($exec*)
}
"""
                with open(os.path.join(cls.YARA_RULES_DIR, "sample_rule.yar"), "w") as f:
                    f.write(sample_rule)
                logger.info("Created sample YARA rule")
            except Exception as e:
                logger.error(f"Error creating YARA rules directory: {e}")
    
    @classmethod
    def update_from_args(cls, args):
        """Update configuration based on command-line arguments"""
        if hasattr(args, 'virustotal') and args.virustotal:
            cls.ENABLED_CHECKS['virustotal'] = True
        
        if hasattr(args, 'no_entropy') and args.no_entropy:
            cls.ENABLED_CHECKS['entropy'] = False
            
        if hasattr(args, 'no_indicators') and args.no_indicators:
            cls.ENABLED_CHECKS['indicators'] = False
            
        if hasattr(args, 'yara') and args.yara:
            cls.ENABLED_CHECKS['yara'] = True
            
        if hasattr(args, 'yara_dir') and args.yara_dir:
            cls.YARA_RULES_DIR = args.yara_dir
            
        if hasattr(args, 'max_size') and args.max_size:
            cls.MAX_FILE_SIZE_MB = args.max_size

# === Advanced Threat Indicators ===
class ThreatIndicators:
    """Comprehensive collection of threat indicators"""
    GENERIC_INDICATORS = [
        # Suspicious command execution patterns
        r'(powershell|cmd\.exe|wscript|mshta|certutil|regsvr32|rundll32)',
        
        # Base64 encoded potential shellcode or commands
        r'(base64,?[A-Za-z0-9+/=]{20,})',
        
        # Dynamic code execution
        r'(eval\(|exec\(|unescape\(|Function\(|window\.eval)',
        
        # Potential obfuscation
        r'(chr\(|fromCharCode\(|atob\()',
        
        # Potential exploitation techniques
        r'(shellcode|bypass|UAC|escalate|exploit)',
    ]

    HTML_SVG_INDICATORS = [
        r'<script.*?>',
        r'on\w+="[^"]+"',
        r'javascript:',
        r'<iframe.*?>',
        r'data:text/html',
        r'window\.location',
    ]

    PDF_INDICATORS = [
        r'/JS', r'/JavaScript', r'/AA', r'/OpenAction',
        r'Launch', r'EmbeddedFile', r'GoToE', 
        r'/URI', r'/GoTo', r'/RichMedia'
    ]

    XML_INDICATORS = [
        r'<!ENTITY',
        r'<!DOCTYPE',
        r'http[s]?://',
        r'file://',
        r'php://filter',
        r'data:',
    ]

# === Threat Intelligence Integration ===
class ThreatIntelligence:
    """Threat intelligence and file reputation checks"""
    @staticmethod
    def virustotal_check(file_path: str) -> Dict[str, Any]:
        """Check file reputation on VirusTotal"""
        if not ScannerConfig.ENABLED_CHECKS['virustotal']:
            return {}
        
        try:
            with open(file_path, 'rb') as f:
                files = {'file': f}
                response = requests.post(
                    'https://www.virustotal.com/vtapi/v2/file/scan',
                    params={'apikey': ScannerConfig.VIRUS_TOTAL_API_KEY},
                    files=files
                )
                
            # Get scan results
            scan_results = response.json()
            
            # Retrieve detailed results
            if 'scan_id' in scan_results:
                results_response = requests.get(
                    'https://www.virustotal.com/vtapi/v2/file/report',
                    params={
                        'apikey': ScannerConfig.VIRUS_TOTAL_API_KEY,
                        'resource': scan_results['scan_id']
                    }
                )
                return results_response.json()
            
        except Exception as e:
            logger.error(f"VirusTotal API error: {e}")
        
        return {}

    @staticmethod
    def hash_file(file_path: str) -> Dict[str, str]:
        """Generate file hashes"""
        hash_algorithms = {
            'md5': hashlib.md5,
            'sha1': hashlib.sha1,
            'sha256': hashlib.sha256
        }
        
        hashes = {}
        with open(file_path, 'rb') as f:
            file_content = f.read()
            for name, algo in hash_algorithms.items():
                hashes[name] = algo(file_content).hexdigest()
        
        return hashes

# === Entropy and Risk Analysis ===
def calculate_entropy(data: bytes) -> float:
    """Calculate Shannon entropy of byte data"""
    if not data:
        return 0.0
    
    # Count byte frequencies
    byte_counts = {}
    for byte in data:
        byte_counts[byte] = byte_counts.get(byte, 0) + 1
    
    # Calculate entropy
    length = len(data)
    entropy = 0.0
    for count in byte_counts.values():
        p = count / length
        entropy -= p * math.log2(p)
    
    return entropy

# Updated YARA rule scanner
class YaraRuleScanner:
    """YARA rule-based malware detection"""
    @staticmethod
    def load_rules(rule_directory: str = None) -> List:
        """Load YARA rules from a directory"""
        rules = []
        
        # Get YARA module
        yara_module = yara_wrapper()
        if yara_module is None:
            logger.warning("Cannot load YARA rules - module unavailable")
            return rules
        
        # Use provided directory or default from config
        if rule_directory is None:
            rule_directory = ScannerConfig.YARA_RULES_DIR
            
        if not os.path.exists(rule_directory):
            logger.warning(f"YARA rule directory {rule_directory} not found.")
            return rules
        
        for filename in os.listdir(rule_directory):
            if filename.endswith('.yar') or filename.endswith('.yara'):
                try:
                    rule_path = os.path.join(rule_directory, filename)
                    
                    # Try compile methods based on what's available
                    if hasattr(yara_module, 'compile'):
                        rule = yara_module.compile(filepath=rule_path)
                        rules.append(rule)
                    elif hasattr(yara_module, 'compile_file'):
                        rule = yara_module.compile_file(filepath=rule_path)
                        rules.append(rule)
                    else:
                        logger.warning(f"No suitable YARA compilation method found for {filename}")
                except Exception as e:
                    logger.error(f"Error compiling YARA rule {filename}: {e}")
        
        return rules

def scan_with_yara_rules(file_path: str, rules: List) -> List[Dict[str, Any]]:
    """Scan a file against loaded YARA rules"""
    if not ScannerConfig.ENABLED_CHECKS['yara']:
        return []
    
    # Try to use the yara_scan_file function from our custom module first
    try:
        matches = yara_scan_file(file_path)
        if matches:
            return [{'rule_name': match.rule, 'tags': match.tags, 'meta': match.meta} 
                   for match in matches]
    except Exception as e:
        logger.warning(f"Could not use yara_scan_file: {e}")
    
    # Fall back to the original implementation if needed
    if not rules:
        return []
        
    matches = []
    for rule_set in rules:
        try:
            # Try different methods to match rules based on yara version
            if hasattr(rule_set, 'match'):
                # Modern yara-python approach
                file_matches = rule_set.match(filepath=file_path)
                
                # Handle different return types
                if isinstance(file_matches, list):
                    for match in file_matches:
                        match_info = {
                            'rule_name': getattr(match, 'rule', 'Unknown'),
                            'tags': getattr(match, 'tags', []),
                            'meta': getattr(match, 'meta', {})
                        }
                        matches.append(match_info)
                else:
                    # Some versions may return a dictionary
                    matches.append({
                        'rule_name': getattr(file_matches, 'rule', 'Unknown'),
                        'tags': getattr(file_matches, 'tags', []),
                        'meta': getattr(file_matches, 'meta', {})
                    })
            else:
                logger.warning(f"Incompatible YARA rule object: {type(rule_set)}")
        except AttributeError as e:
            logger.error(f"YARA attribute error: {e}")
        except Exception as e:
            logger.error(f"YARA scanning error: {e}")
    
    return matches

# === Content Scanning Functions ===
def scan_content(content: str, indicators: List[str], label: str = "Generic") -> List[Tuple[str, str, List[str]]]:
    """Scan content against threat indicators"""
    findings = []
    for pattern in indicators:
        matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
        if matches:
            findings.append((label, pattern, matches[:5]))
    return findings

def scan_file(file_path: str) -> Dict[str, Any]:
    """Comprehensive file scanning"""
    # File size check
    file_size = os.path.getsize(file_path)
    if file_size > ScannerConfig.MAX_FILE_SIZE_MB * 1024 * 1024:
        logger.warning(f"File {file_path} exceeds maximum scan size.")
        return {}
    
    # Scan results container
    scan_results = {
        'file_path': file_path,
        'file_size': file_size,
        'findings': [],
        'hashes': {},
        'entropy': 0.0,
        'virustotal': {},
        'yara_matches': []
    }
    
    # Compute file hashes
    scan_results['hashes'] = ThreatIntelligence.hash_file(file_path)
    
    # Read entire file content
    with open(file_path, 'rb') as f:
        file_content = f.read()
    
    # Entropy check
    if ScannerConfig.ENABLED_CHECKS['entropy']:
        entropy = calculate_entropy(file_content)
        scan_results['entropy'] = entropy
        logger.info(f"File entropy: {entropy:.2f}")
        if entropy > ScannerConfig.ENTROPY_THRESHOLD:
            scan_results['findings'].append(
                ("Entropy", "High entropy", [f"Entropy: {entropy:.2f}"])
            )
    
    # VirusTotal check
    if ScannerConfig.ENABLED_CHECKS['virustotal']:
        scan_results['virustotal'] = ThreatIntelligence.virustotal_check(file_path)
    
    # YARA rule scanning - use the improved individual rule scanning approach
    if ScannerConfig.ENABLED_CHECKS['yara']:
        try:
            # Import from malscan_yara instead of yara
            from malscan_yara import scan_with_individual_rules
            
            # First try to use the simplified rules (more reliable)
            simple_rules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'simple_yara_rules')
            if os.path.exists(simple_rules_dir):
                yara_matches = scan_with_individual_rules(file_path, rules_dir=simple_rules_dir)
                scan_results['yara_matches'].extend(yara_matches)
                
            # Then try the main rule set if enabled and we didn't find anything yet
            if not scan_results['yara_matches'] or ScannerConfig.ENABLED_CHECKS.get('all_yara_rules', False):
                matches = scan_with_individual_rules(file_path, rules_dir=ScannerConfig.YARA_RULES_DIR)
                scan_results['yara_matches'].extend(matches)
                
            logger.info(f"YARA matches: {len(scan_results['yara_matches'])}")
        except Exception as e:
            logger.error(f"YARA scanning error: {e}")
            scan_results['yara_matches'] = []
    
    # File type-specific scanning
    ftype = detect_filetype(file_path)
    logger.info(f"Detected file type: {ftype}")
    
    try:
        if ftype == "pdf":
            logger.info(f"Scanning as PDF file")
            try:
                # Extract text from PDF
                pdf_text = extract_text(file_path)
                
                # Scan PDF text content
                findings = scan_content(
                    pdf_text, 
                    ThreatIndicators.PDF_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "PDF"
                )
                scan_results['findings'] += findings
            except Exception as e:
                logger.error(f"Error in PDF scanning: {e}")
                
        elif ftype in ["html", "svg"]:
            logger.info(f"Scanning as HTML/SVG file")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                soup = BeautifulSoup(content, "html.parser")
                raw = soup.prettify()
                
                # Debug output
                logger.info(f"File content length: {len(raw)} characters")
                
                # Perform the scan
                findings = scan_content(
                    raw, 
                    ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "HTML/SVG"
                )
                
                # Debug output
                logger.info(f"Found {len(findings)} indicators in HTML/SVG content")
                
                scan_results['findings'] += findings
            except Exception as e:
                logger.error(f"Error in HTML/SVG scanning: {e}")
                
        elif ftype == "xml":
            logger.info(f"Scanning as XML file")
            try:
                # Use proper XML parser
                tree = ET.parse(file_path)
                root = tree.getroot()
                # Convert to string for scanning
                content = ET.tostring(root, encoding='unicode')
                scan_results['findings'] += scan_content(
                    content, 
                    ThreatIndicators.XML_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "XML"
                )
            except ET.ParseError as e:
                logger.warning(f"Could not parse {file_path} as XML: {e}")
                # Fallback to text scanning
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                scan_results['findings'] += scan_content(
                    content, 
                    ThreatIndicators.XML_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "XML"
                )
                
        elif ftype == "docx":
            logger.info(f"Scanning as DOCX file")
            try:
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                scan_results['findings'] += scan_content(
                    text,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "DOCX"
                )
            except Exception as e:
                logger.error(f"Error in DOCX scanning: {e}")
                
        elif ftype == "xlsx":
            logger.info(f"Scanning as XLSX file")
            try:
                wb = openpyxl.load_workbook(file_path)
                text = []
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True):
                        text.append(" ".join([str(cell) for cell in row if cell]))
                scan_results['findings'] += scan_content(
                    "\n".join(text),
                    ThreatIndicators.GENERIC_INDICATORS,
                    "XLSX"
                )
            except Exception as e:
                logger.error(f"Error in XLSX scanning: {e}")
                
        elif ftype == "zip":
            logger.info(f"Scanning as ZIP archive")
            try:
                with zipfile.ZipFile(file_path) as zf:
                    # Check for suspicious files in the archive
                    suspicious_extensions = ['.exe', '.dll', '.js', '.vbs', '.hta', '.ps1']
                    suspicious_files = []
                    
                    for name in zf.namelist():
                        ext = os.path.splitext(name)[1].lower()
                        if ext in suspicious_extensions:
                            suspicious_files.append(name)
                    
                    if suspicious_files:
                        scan_results['findings'].append(
                            ("ZIP", "Suspicious files in archive", suspicious_files[:5])
                        )
                    
                    # Scan selected files within the archive
                    with tempfile.TemporaryDirectory() as temp_dir:
                        for name in zf.namelist():
                            ext = os.path.splitext(name)[1].lower()
                            if ext in ['.html', '.xml', '.svg', '.pdf', '.js', '.txt']:
                                zf.extract(name, temp_dir)
                                extracted_path = os.path.join(temp_dir, name)
                                
                                # Only scan if file size is reasonable
                                if os.path.getsize(extracted_path) < ScannerConfig.MAX_FILE_SIZE_MB * 1024 * 1024:
                                    # Scan the extracted file
                                    sub_results = scan_file(extracted_path)
                                    
                                    # Add findings to main results
                                    if sub_results.get('findings'):
                                        for finding in sub_results['findings']:
                                            scan_results['findings'].append(
                                                (f"ZIP/{name}", finding[1], finding[2])
                                            )
            except Exception as e:
                logger.error(f"Error in ZIP scanning: {e}")
                
        elif ftype == "eml":
            logger.info(f"Scanning as email file")
            try:
                with open(file_path, 'rb') as email_file:
                    msg = BytesParser(policy=policy.default).parse(email_file)
                
                # Scan email headers
                headers = ""
                for header, value in msg.items():
                    headers += f"{header}: {value}\n"
                
                scan_results['findings'] += scan_content(
                    headers,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Email Headers"
                )
                
                # Scan email body
                if msg.is_multipart():
                    for part in msg.iter_parts():
                        content_type = part.get_content_type()
                        if content_type == 'text/plain':
                            text = part.get_content()
                            scan_results['findings'] += scan_content(
                                text,
                                ThreatIndicators.GENERIC_INDICATORS,
                                "Email Body (plain)"
                            )
                        elif content_type == 'text/html':
                            html = part.get_content()
                            scan_results['findings'] += scan_content(
                                html,
                                ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS,
                                "Email Body (HTML)"
                            )
                else:
                    # Not multipart - get content directly
                    content = msg.get_content()
                    content_type = msg.get_content_type()
                    
                    if content_type == 'text/plain':
                        scan_results['findings'] += scan_content(
                            content,
                            ThreatIndicators.GENERIC_INDICATORS,
                            "Email Body (plain)"
                        )
                    elif content_type == 'text/html':
                        scan_results['findings'] += scan_content(
                            content,
                            ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS,
                            "Email Body (HTML)"
                        )
            except Exception as e:
                logger.error(f"Error in email scanning: {e}")
                
        elif ftype == "text":
            logger.info(f"Scanning as text file")
            try:
                # Detect encoding
                with open(file_path, 'rb') as f:
                    raw_data = f.read(4096)  # Read a chunk to detect encoding
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                
                # Read file with detected encoding
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                
                scan_results['findings'] += scan_content(
                    content,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Text"
                )
            except Exception as e:
                logger.error(f"Error in text file scanning: {e}")
                
        else:
            logger.info(f"Generic binary scanning for {ftype}")
            # For unknown file types, scan for embedded strings
            try:
                # Extract strings from binary
                strings_output = []
                current_string = ""
                
                for byte in file_content:
                    if 32 <= byte <= 126:  # Printable ASCII
                        current_string += chr(byte)
                    else:
                        if len(current_string) >= 4:  # Only keep strings of reasonable length
                            strings_output.append(current_string)
                        current_string = ""
                
                # Add the last string if it exists
                if len(current_string) >= 4:
                    strings_output.append(current_string)
                
                # Join strings and scan
                strings_text = "\n".join(strings_output)
                scan_results['findings'] += scan_content(
                    strings_text,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Binary Strings"
                )
            except Exception as e:
                logger.error(f"Error in binary string extraction: {e}")
    
    except Exception as e:
        logger.error(f"Error scanning {file_path}: {e}")
    
    # Debug output
    logger.info(f"Total findings: {len(scan_results['findings'])}")
    if len(scan_results['findings']) > 0:
        logger.info(f"First finding: {scan_results['findings'][0]}")
    
    return scan_results

def detect_filetype(filepath: str) -> str:
    """Improved file type detection"""
    ext = os.path.splitext(filepath)[1].lower()
    
    # Mapping of extensions to specific types
    type_map = {
        ".pdf": "pdf",
        ".svg": "svg",
        ".html": "html",
        ".htm": "html",
        ".xml": "xml",
        ".zip": "zip",
        ".txt": "text",
        ".csv": "text",
        ".log": "text",
        ".json": "text",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".eml": "eml"
    }
    
    # Check exact extension
    if ext in type_map:
        return type_map[ext]
    
    # Fallback to MIME type detection
    mime = mimetypes.guess_type(filepath)[0]
    if mime:
        if mime.startswith("text"):
            return "text"
        if "xml" in mime:
            return "xml"
        if "html" in mime:
            return "html"
    
    return "unknown"

def export_findings(scan_results: Dict[str, Any], output_prefix: str):
    """Export scan results to multiple formats"""
    csv_path = f"{output_prefix}_scan_results.csv"
    json_path = f"{output_prefix}_scan_results.json"
    
    # Consolidate findings
    all_findings = []
    
    # Add generic findings
    for label, pattern, matches in scan_results.get('findings', []):
        for match in matches:
            all_findings.append({
                'type': label,
                'pattern': pattern,
                'match': match
            })
    
    # Add YARA matches
    for yara_match in scan_results.get('yara_matches', []):
        all_findings.append({
            'type': 'YARA',
            'rule_name': yara_match.get('rule_name', ''),
            'tags': yara_match.get('tags', []),
            'meta': yara_match.get('meta', {})
        })
    
    # Write CSV
    with open(csv_path, "w", newline="", encoding="utf-8") as csvfile:
        fieldnames = ['type', 'pattern', 'match']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for finding in all_findings:
            writer.writerow(finding)
    
    # Write JSON
    with open(json_path, "w", encoding="utf-8") as jsonfile:
        json.dump({
            'scan_results': scan_results,
            'findings': all_findings
        }, jsonfile, indent=2)
    
    logger.info(f"Scan results exported to {csv_path} and {json_path}")

def scan_directory(directory: str, recursive: bool = False):
    """Scan an entire directory for potential threats"""
    scan_results = []
    
    def scan_files(dir_path):
        nonlocal scan_results
        for root, _, files in os.walk(dir_path):
            for filename in files:
                filepath = os.path.join(root, filename)
                try:
                    result = scan_file(filepath)
                    if result.get('findings') or result.get('yara_matches'):
                        scan_results.append(result)
                except Exception as e:
                    logger.error(f"Error scanning {filepath}: {e}")
                
                # Break early if recursive is False
                if not recursive:
                    break
    
    # Start scanning
    scan_files(directory)
    return scan_results

def main():
    """Main CLI entry point for the malware scanner"""
    parser = argparse.ArgumentParser(description="Advanced Malware Scanner")
    
    # Create default directories
    ScannerConfig.create_default_directories()
    
    # Initialize YARA functionality
    from malscan_yara import yara_manager
    if not yara_manager.is_yara_available():
        logger.warning("YARA functionality limited - ensure yara-python is installed")
    else:
        logger.info("YARA functionality available and initialized")
    
    # Create simple rules if needed
    simple_rules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'simple_yara_rules')
    if not os.path.exists(simple_rules_dir):
        logger.info("Creating simple YARA rules...")
        create_simple_yara_rules(simple_rules_dir)
    
    # Add mutually exclusive group for file/directory scanning
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("-f", "--file", help="Path to the file to scan")
    group.add_argument("-d", "--directory", help="Directory to scan")
    
    # Additional optional arguments
    parser.add_argument("--recursive", action="store_true", 
                        help="Recursively scan subdirectories")
    parser.add_argument("--export", action="store_true", 
                        help="Export scan results to CSV and JSON")
    parser.add_argument("--virustotal", action="store_true", 
                        help="Enable VirusTotal scanning (requires API key)")
    parser.add_argument("--yara", action="store_true", 
                        help="Enable YARA rule scanning")
    parser.add_argument("--all-yara-rules", action="store_true", 
                        help="Use all YARA rules (can be slow with community rules)")
    parser.add_argument("--no-entropy", action="store_true", 
                        help="Disable entropy scanning")
    parser.add_argument("--no-indicators", action="store_true", 
                        help="Disable threat indicator scanning")
    parser.add_argument("--yara-dir", 
                        help="Custom YARA rules directory")
    parser.add_argument("--max-size", type=int, 
                        help="Maximum file size in MB to scan")
    parser.add_argument("--output", help="Output directory for scan results")
    
    # Parse arguments
    args = parser.parse_args()
    
    # Configure scanner based on CLI arguments
    ScannerConfig.update_from_args(args)
    
    # Add custom configuration for all YARA rules
    if hasattr(args, 'all_yara_rules') and args.all_yara_rules:
        ScannerConfig.ENABLED_CHECKS['all_yara_rules'] = True
    
    # Set output directory
    output_dir = args.output if hasattr(args, 'output') and args.output else os.getcwd()
    
    # Perform scanning
    if args.file:
        if not os.path.isfile(args.file):
            logger.error(f"File not found: {args.file}")
            return
        
        logger.info(f"Scanning file: {args.file}")
        result = scan_file(args.file)
        
        # Print and optionally export results
        print_findings(result)
        
        if args.export:
            output_prefix = os.path.join(output_dir, os.path.splitext(os.path.basename(args.file))[0])
            export_findings(result, output_prefix)
    
    elif args.directory:
        if not os.path.isdir(args.directory):
            logger.error(f"Directory not found: {args.directory}")
            return
        
        logger.info(f"Scanning directory: {args.directory}")
        results = scan_directory(args.directory, recursive=args.recursive)
        
        # Summary of findings
        total_files_scanned = len(results)
        files_with_findings = sum(1 for r in results if r.get('findings') or r.get('yara_matches'))
        
        print(f"\n[SCAN SUMMARY]")
        print(f"Total files scanned: {total_files_scanned}")
        print(f"Files with potential threats: {files_with_findings}")
        
        # Print detailed findings
        for result in results:
            if result.get('findings') or result.get('yara_matches'):
                print(f"\n--- Scan Results for {result['file_path']} ---")
                print_findings(result)
        
        # Export if requested
        if args.export:
            output_prefix = os.path.join(output_dir, os.path.basename(args.directory))
            export_findings(
                {'scanned_files': results}, 
                f"{output_prefix}_scan_results"
            )

def create_simple_yara_rules(output_dir):
    """Create a set of simple, reliable YARA rules for basic scanning"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Basic malware detection rule
    basic_rule = """
rule SuspiciousBehavior {
    meta:
        description = "Detects common suspicious patterns"
        author = "Security Scanner"
        severity = "medium"
    
    strings:
        $shell1 = "cmd.exe" nocase
        $shell2 = "powershell" nocase
        $shell3 = "bash" nocase
        $shell4 = "/bin/sh" nocase
        $exec1 = "exec(" nocase
        $exec2 = "eval(" nocase
        $exec3 = "system(" nocase
        $obf1 = "base64" nocase
        $obf2 = "encode" nocase
        $net1 = "http://" nocase
        $net2 = "https://" nocase
        $sus1 = "exploit" nocase
        $sus2 = "bypass" nocase
        $sus3 = "inject" nocase
    
    condition:
        (2 of ($shell*)) or
        (any of ($exec*)) or
        (any of ($obf*) and any of ($net*)) or
        (2 of ($sus*))
}
"""
    
    # SVG specific rule
    svg_rule = """
rule SuspiciousSVG {
    meta:
        description = "Detects suspicious patterns in SVG files"
        author = "Security Scanner"
        severity = "high"
    
    strings:
        $script_tag = /<script[^>]*>.*<\/script>/s nocase
        $event_handler = /on\w+\s*=/i
        $iframe = /<iframe/i
        $eval = /eval\s*\(/i
        $js_uri = /javascript:/i
        $data_uri = /data:text\/html/i
        $foreignObject = /<foreignObject/i
        $embed = /<embed/i
    
    condition:
        any of them
}
"""
    
    # Document macro rule
    macro_rule = """
rule SuspiciousMacro {
    meta:
        description = "Detects suspicious patterns in document macros and scripts"
        author = "Security Scanner"
        severity = "high"
    
    strings:
        $auto_open = "Auto_Open" nocase
        $auto_exec = "AutoExec" nocase
        $auto_exit = "Auto_Exit" nocase
        $auto_close = "Auto_Close" nocase
        $document_open = "Document_Open" nocase
        $workbook_open = "Workbook_Open" nocase
        
        $shell = "Shell" nocase
        $wscript = "WScript" nocase
        $powershell = "powershell" nocase
        
        $create_object = "CreateObject" nocase
        $get_object = "GetObject" nocase
        
        $http_download = "DownloadFile" nocase
        $xml_http = "XMLHTTP" nocase
        $adodb = "ADODB.Stream" nocase
        
    condition:
        (any of ($auto_*) or $document_open or $workbook_open) and
        (
            any of ($shell, $wscript, $powershell) or
            any of ($create_object, $get_object) or
            any of ($http_download, $xml_http, $adodb)
        )
}
"""
    
    # Write rules to files
    with open(os.path.join(output_dir, "basic.yar"), "w") as f:
        f.write(basic_rule)
        
    with open(os.path.join(output_dir, "svg.yar"), "w") as f:
        f.write(svg_rule)
        
    with open(os.path.join(output_dir, "macro.yar"), "w") as f:
        f.write(macro_rule)
    
    logger.info(f"Created simple YARA rules in {output_dir}")
    return output_dir

def scan_file(file_path: str) -> Dict[str, Any]:
    """Comprehensive file scanning"""
    # File size check
    file_size = os.path.getsize(file_path)
    if file_size > ScannerConfig.MAX_FILE_SIZE_MB * 1024 * 1024:
        logger.warning(f"File {file_path} exceeds maximum scan size.")
        return {}
    
    # Scan results container
    scan_results = {
        'file_path': file_path,
        'file_size': file_size,
        'findings': [],
        'hashes': {},
        'entropy': 0.0,
        'virustotal': {},
        'yara_matches': []
    }
    
    # Compute file hashes
    scan_results['hashes'] = ThreatIntelligence.hash_file(file_path)
    
    # Read entire file content
    with open(file_path, 'rb') as f:
        file_content = f.read()
    
    # Entropy check
    if ScannerConfig.ENABLED_CHECKS['entropy']:
        entropy = calculate_entropy(file_content)
        scan_results['entropy'] = entropy
        logger.info(f"File entropy: {entropy:.2f}")
        if entropy > ScannerConfig.ENTROPY_THRESHOLD:
            scan_results['findings'].append(
                ("Entropy", "High entropy", [f"Entropy: {entropy:.2f}"])
            )
    
    # VirusTotal check
    if ScannerConfig.ENABLED_CHECKS['virustotal']:
        scan_results['virustotal'] = ThreatIntelligence.virustotal_check(file_path)
    
    # YARA rule scanning - use the improved individual rule scanning approach with error suppression
    if ScannerConfig.ENABLED_CHECKS['yara']:
        try:
            # Temporarily suppress logging for YARA operations
            yara_logger = logging.getLogger('root')
            original_level = yara_logger.level
            yara_logger.setLevel(logging.CRITICAL)  # Only critical errors
            
            try:
                # Import from malscan_yara instead of yara
                from malscan_yara import scan_with_individual_rules
                
                # First try to use the simplified rules (more reliable)
                simple_rules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'simple_yara_rules')
                if os.path.exists(simple_rules_dir):
                    yara_matches = scan_with_individual_rules(file_path, rules_dir=simple_rules_dir)
                    if yara_matches:
                        scan_results['yara_matches'].extend(yara_matches)
                        logger.info(f"Found {len(yara_matches)} matches in simple YARA rules")
                    
                # Then try the main rule set if enabled and we didn't find anything yet
                if not scan_results['yara_matches'] or ScannerConfig.ENABLED_CHECKS.get('all_yara_rules', False):
                    matches = scan_with_individual_rules(file_path, rules_dir=ScannerConfig.YARA_RULES_DIR)
                    if matches:
                        scan_results['yara_matches'].extend(matches)
                        logger.info(f"Found {len(matches)} matches in full YARA ruleset")
                
            finally:
                # Restore original logging level
                yara_logger.setLevel(original_level)
                
            # Only log if we actually found something
            if scan_results['yara_matches']:
                logger.info(f"Total YARA matches: {len(scan_results['yara_matches'])}")
                
        except Exception:
            # Silently continue without error messages
            pass
    
    # File type-specific scanning
    ftype = detect_filetype(file_path)
    logger.info(f"Detected file type: {ftype}")
    
    try:
        if ftype == "pdf":
            logger.info(f"Scanning as PDF file")
            try:
                # Extract text from PDF
                pdf_text = extract_text(file_path)
                
                # Scan PDF text content
                findings = scan_content(
                    pdf_text, 
                    ThreatIndicators.PDF_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "PDF"
                )
                scan_results['findings'] += findings
            except Exception as e:
                logger.error(f"Error in PDF scanning: {e}")
                
        elif ftype in ["html", "svg"]:
            logger.info(f"Scanning as HTML/SVG file")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                soup = BeautifulSoup(content, "html.parser")
                raw = soup.prettify()
                
                # Debug output
                logger.info(f"File content length: {len(raw)} characters")
                
                # Perform the scan
                findings = scan_content(
                    raw, 
                    ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "HTML/SVG"
                )
                
                # Debug output
                logger.info(f"Found {len(findings)} indicators in HTML/SVG content")
                
                scan_results['findings'] += findings
            except Exception as e:
                logger.error(f"Error in HTML/SVG scanning: {e}")
                
        elif ftype == "xml":
            logger.info(f"Scanning as XML file")
            try:
                # Use proper XML parser
                tree = ET.parse(file_path)
                root = tree.getroot()
                # Convert to string for scanning
                content = ET.tostring(root, encoding='unicode')
                scan_results['findings'] += scan_content(
                    content, 
                    ThreatIndicators.XML_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "XML"
                )
            except ET.ParseError as e:
                logger.warning(f"Could not parse {file_path} as XML: {e}")
                # Fallback to text scanning
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                scan_results['findings'] += scan_content(
                    content, 
                    ThreatIndicators.XML_INDICATORS + ThreatIndicators.GENERIC_INDICATORS, 
                    "XML"
                )
                
        elif ftype == "docx":
            logger.info(f"Scanning as DOCX file")
            try:
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                scan_results['findings'] += scan_content(
                    text,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "DOCX"
                )
            except Exception as e:
                logger.error(f"Error in DOCX scanning: {e}")
                
        elif ftype == "xlsx":
            logger.info(f"Scanning as XLSX file")
            try:
                wb = openpyxl.load_workbook(file_path)
                text = []
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True):
                        text.append(" ".join([str(cell) for cell in row if cell]))
                scan_results['findings'] += scan_content(
                    "\n".join(text),
                    ThreatIndicators.GENERIC_INDICATORS,
                    "XLSX"
                )
            except Exception as e:
                logger.error(f"Error in XLSX scanning: {e}")
                
        elif ftype == "zip":
            logger.info(f"Scanning as ZIP archive")
            try:
                with zipfile.ZipFile(file_path) as zf:
                    # Check for suspicious files in the archive
                    suspicious_extensions = ['.exe', '.dll', '.js', '.vbs', '.hta', '.ps1']
                    suspicious_files = []
                    
                    for name in zf.namelist():
                        ext = os.path.splitext(name)[1].lower()
                        if ext in suspicious_extensions:
                            suspicious_files.append(name)
                    
                    if suspicious_files:
                        scan_results['findings'].append(
                            ("ZIP", "Suspicious files in archive", suspicious_files[:5])
                        )
                    
                    # Scan selected files within the archive
                    with tempfile.TemporaryDirectory() as temp_dir:
                        for name in zf.namelist():
                            ext = os.path.splitext(name)[1].lower()
                            if ext in ['.html', '.xml', '.svg', '.pdf', '.js', '.txt']:
                                zf.extract(name, temp_dir)
                                extracted_path = os.path.join(temp_dir, name)
                                
                                # Only scan if file size is reasonable
                                if os.path.getsize(extracted_path) < ScannerConfig.MAX_FILE_SIZE_MB * 1024 * 1024:
                                    # Scan the extracted file
                                    sub_results = scan_file(extracted_path)
                                    
                                    # Add findings to main results
                                    if sub_results.get('findings'):
                                        for finding in sub_results['findings']:
                                            scan_results['findings'].append(
                                                (f"ZIP/{name}", finding[1], finding[2])
                                            )
            except Exception as e:
                logger.error(f"Error in ZIP scanning: {e}")
                
        elif ftype == "eml":
            logger.info(f"Scanning as email file")
            try:
                with open(file_path, 'rb') as email_file:
                    msg = BytesParser(policy=policy.default).parse(email_file)
                
                # Scan email headers
                headers = ""
                for header, value in msg.items():
                    headers += f"{header}: {value}\n"
                
                scan_results['findings'] += scan_content(
                    headers,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Email Headers"
                )
                
                # Scan email body
                if msg.is_multipart():
                    for part in msg.iter_parts():
                        content_type = part.get_content_type()
                        if content_type == 'text/plain':
                            text = part.get_content()
                            scan_results['findings'] += scan_content(
                                text,
                                ThreatIndicators.GENERIC_INDICATORS,
                                "Email Body (plain)"
                            )
                        elif content_type == 'text/html':
                            html = part.get_content()
                            scan_results['findings'] += scan_content(
                                html,
                                ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS,
                                "Email Body (HTML)"
                            )
                else:
                    # Not multipart - get content directly
                    content = msg.get_content()
                    content_type = msg.get_content_type()
                    
                    if content_type == 'text/plain':
                        scan_results['findings'] += scan_content(
                            content,
                            ThreatIndicators.GENERIC_INDICATORS,
                            "Email Body (plain)"
                        )
                    elif content_type == 'text/html':
                        scan_results['findings'] += scan_content(
                            content,
                            ThreatIndicators.HTML_SVG_INDICATORS + ThreatIndicators.GENERIC_INDICATORS,
                            "Email Body (HTML)"
                        )
            except Exception as e:
                logger.error(f"Error in email scanning: {e}")
                
        elif ftype == "text":
            logger.info(f"Scanning as text file")
            try:
                # Detect encoding
                with open(file_path, 'rb') as f:
                    raw_data = f.read(4096)  # Read a chunk to detect encoding
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                
                # Read file with detected encoding
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                
                scan_results['findings'] += scan_content(
                    content,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Text"
                )
            except Exception as e:
                logger.error(f"Error in text file scanning: {e}")
                
        else:
            logger.info(f"Generic binary scanning for {ftype}")
            # For unknown file types, scan for embedded strings
            try:
                # Extract strings from binary
                strings_output = []
                current_string = ""
                
                for byte in file_content:
                    if 32 <= byte <= 126:  # Printable ASCII
                        current_string += chr(byte)
                    else:
                        if len(current_string) >= 4:  # Only keep strings of reasonable length
                            strings_output.append(current_string)
                        current_string = ""
                
                # Add the last string if it exists
                if len(current_string) >= 4:
                    strings_output.append(current_string)
                
                # Join strings and scan
                strings_text = "\n".join(strings_output)
                scan_results['findings'] += scan_content(
                    strings_text,
                    ThreatIndicators.GENERIC_INDICATORS,
                    "Binary Strings"
                )
            except Exception as e:
                logger.error(f"Error in binary string extraction: {e}")
    
    except Exception as e:
        logger.error(f"Error scanning {file_path}: {e}")
    
    # Debug output
    logger.info(f"Total findings: {len(scan_results['findings'])}")
    if len(scan_results['findings']) > 0:
        logger.info(f"First finding: {scan_results['findings'][0]}")
    
    return scan_results

def yara_wrapper():
    """Wrapper to manage YARA operations and handle missing functionality"""
    try:
        import yara
        
        # Test if yara is properly installed
        if not hasattr(yara, 'compile') and not hasattr(yara, 'compile_file'):
            logger.warning("YARA module installed but missing required methods")
            return None
            
        return yara
    except ImportError:
        logger.warning("YARA module not installed")
        return None
        
def print_findings(scan_result: Dict[str, Any]):
    """Pretty print scan findings using the detailed threat report"""
    if not scan_result:
        logger.info("No scan results to display.")
        return
    
    # Generate comprehensive threat report
    file_info = {
        'size': scan_result.get('file_size', 0),
        'file_type': detect_filetype(scan_result.get('file_path', '')),
        'entropy': scan_result.get('entropy', 0),
        'md5': scan_result.get('hashes', {}).get('md5', 'Unknown'),
        'sha1': scan_result.get('hashes', {}).get('sha1', 'Unknown'),
        'sha256': scan_result.get('hashes', {}).get('sha256', 'Unknown'),
        'yara_matches': scan_result.get('yara_matches', [])
    }
    
    findings = scan_result.get('findings', [])
    
    # Generate and print the threat report
    report = generate_threat_report(
        scan_result.get('file_path', 'Unknown File'),
        findings,
        file_info
    )
    
    print(report)
    
    # VirusTotal results - only if there are positives (not included in threat report)
    vt_results = scan_result.get('virustotal', {})
    if vt_results and vt_results.get('positives', 0) > 0:
        print("\n[VirusTotal Scan]")
        print(f"Positives: {vt_results['positives']}/{vt_results.get('total', 'N/A')}")
        print(f"Scan Date: {vt_results.get('scan_date', 'N/A')}")

def generate_threat_report(file_path, findings, file_info):
    """
    Generate a comprehensive threat report with detailed explanations of findings.
    
    Args:
        file_path (str): Path to the analyzed file
        findings (list): List of threat indicator findings
        file_info (dict): File metadata including hashes and entropy
        
    Returns:
        str: Formatted threat report with detailed explanations
    """
    report = []
    
    # File header section
    report.append(f"=== THREAT ANALYSIS REPORT ===")
    report.append(f"File: {os.path.basename(file_path)}")
    report.append(f"Analysis Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append("")
    
    # File information section
    report.append("[File Information]")
    report.append(f"Size: {file_info.get('size', 'Unknown')} bytes")
    report.append(f"Type: {file_info.get('file_type', 'Unknown')}")
    report.append(f"Entropy: {file_info.get('entropy', 'Unknown')} ({interpret_entropy(file_info.get('entropy', 0))})")
    report.append("")
    
    # File hashes section
    report.append("[File Hashes]")
    report.append(f"MD5: {file_info.get('md5', 'Unknown')}")
    report.append(f"SHA1: {file_info.get('sha1', 'Unknown')}")
    report.append(f"SHA256: {file_info.get('sha256', 'Unknown')}")
    report.append("")
    
    # YARA matches section (if any)
    if 'yara_matches' in file_info and file_info['yara_matches']:
        report.append("[YARA Rule Matches]")
        for match in file_info['yara_matches']:
            # Check if match is a dictionary or an object and handle accordingly
            if isinstance(match, dict):
                # Match is a dictionary
                rule_name = match.get('rule_name', 'Unknown')
                namespace = match.get('namespace', 'Unknown')
                tags = match.get('tags', [])
                meta = match.get('meta', {})
                
                report.append(f"Rule: {rule_name} (Namespace: {namespace})")
                
                # Add tags if available
                if tags:
                    report.append(f"  Tags: {', '.join(tags)}")
                    
                # Add metadata if available
                if meta:
                    report.append("  Metadata:")
                    for key, value in meta.items():
                        report.append(f"    {key}: {value}")
                
                # Add strings if available
                if 'strings' in match and match['strings']:
                    report.append("  Matched strings:")
                    for string in match['strings']:
                        identifier = string.get('identifier', 'Unknown')
                        offset = string.get('offset', 0)
                        data = string.get('data', '')
                        report.append(f"    - {identifier} at offset {offset}: {data}")
            else:
                # Match is an object with attributes (original code)
                try:
                    rule_name = getattr(match, 'rule', 'Unknown')
                    namespace = getattr(match, 'namespace', 'Unknown')
                    report.append(f"Rule: {rule_name} (Namespace: {namespace})")
                    
                    if hasattr(match, 'strings') and match.strings:
                        report.append("  Matched strings:")
                        for string in match.strings:
                            report.append(f"    - {string.identifier} at offset {string.offset}: {repr(string.data[:100])}")
                except AttributeError as e:
                    # Fallback for any unexpected object structure
                    report.append(f"Rule: {str(match)}")
                    report.append(f"  (Could not extract detailed information: {e})")
            
            report.append("")
    
    # Threat indicators section with detailed explanations
    if findings:
        report.append("[Threat Indicators]")
        
        # Use a set to track unique findings to avoid repetition
        processed_matches = set()
        
        for finding in findings:
            indicator_type, pattern, matches = finding
            for match in matches:
                # Skip if we've already processed this exact match
                match_key = f"{pattern}:{match}"
                if match_key in processed_matches:
                    continue
                processed_matches.add(match_key)
                
                # Get detailed explanation
                explanation = explain_threat_indicator(indicator_type, pattern, match)
                report.append(explanation)
        
        # Summary section
        risk_level = calculate_risk_level(findings)
        report.append("")
        report.append("[Risk Assessment]")
        report.append(f"Overall Risk Level: {risk_level}")
        report.append(f"Total Indicators Found: {sum(len(matches) for _, _, matches in findings)}")
        report.append("")
    else:
        report.append("[Threat Indicators]")
        report.append("No threat indicators found in this file.")
        report.append("")
        report.append("[Risk Assessment]")
        report.append("Overall Risk Level: LOW")
        report.append("")
    
    # Recommendations section
    report.append("[Recommendations]")
    if findings:
        report.append("1. Isolate this file and do not open it in browsers or SVG renderers")
        report.append("2. Review the identified threat indicators in context")
        report.append("3. Use a sandbox environment for further analysis if needed")
        report.append("4. Consider scanning with additional tools for confirmation")
    else:
        report.append("No suspicious indicators were found, but always follow security best practices:")
        report.append("1. Only use files from trusted sources")
        report.append("2. Keep security software up to date")
    
    return "\n".join(report)

def explain_threat_indicator(indicator_type, pattern, match):
    """
    Provide detailed explanation for threat indicators found in files.
    
    Args:
        indicator_type (str): Type of indicator (e.g., 'HTML/SVG')
        pattern (str): Pattern that was matched
        match (str): Actual content that matched the pattern
        
    Returns:
        str: Detailed explanation of the indicator and its security implications
    """
    explanations = {
        '<script.*?>': {
            'title': 'Embedded Script Tag',
            'description': 'Script tags in SVG files can execute arbitrary JavaScript code when the image is loaded in a browser. This is a common vector for cross-site scripting (XSS) attacks and malware delivery.',
            'impact': 'HIGH',
            'mitre_technique': 'T1059.007 (Command and Scripting Interpreter: JavaScript)',
            'remediation': 'Remove script tags from SVG files or sanitize them before rendering.'
        },
        'on\\w+="[^"]+"': {
            'title': 'Event Handler Attribute',
            'description': 'Event handlers (like onclick, onload) in SVG files can execute JavaScript when specific events occur. Attackers use these to trigger malicious code without requiring explicit script tags.',
            'impact': 'MEDIUM',
            'mitre_technique': 'T1059.007 (Command and Scripting Interpreter: JavaScript)',
            'remediation': 'Remove event handler attributes from SVG elements.'
        },
        '(eval\\(|exec\\(|unescape\\(|Function\\(|window\\.eval)': {
            'title': 'Dynamic Code Execution Function',
            'description': 'Functions like eval() execute JavaScript code from strings. This is commonly used to obfuscate malicious code and evade detection. It can execute encoded or encrypted payloads.',
            'impact': 'CRITICAL',
            'mitre_technique': 'T1027 (Obfuscated Files or Information)',
            'remediation': 'Remove dynamic execution functions as they are rarely needed in legitimate SVG files.'
        },
        '(chr\\(|fromCharCode\\(|atob\\()': {
            'title': 'String Encoding/Decoding Function',
            'description': 'Functions like atob() decode Base64 encoded strings. Used to hide malicious payloads from basic detection. Often paired with eval() to execute the decoded content.',
            'impact': 'HIGH',
            'mitre_technique': 'T1027.010 (Obfuscated Files or Information: Command Obfuscation)',
            'remediation': 'Remove encoding/decoding functions and investigate any encoded content.'
        },
        'base64': {
            'title': 'Base64 Encoded Content',
            'description': 'Base64 encoding converts binary data to ASCII text. Malicious actors often use it to hide payloads, command and control URLs, or exfiltrated data.',
            'impact': 'MEDIUM',
            'mitre_technique': 'T1132.001 (Data Encoding: Standard Encoding)',
            'remediation': 'Decode and analyze any Base64 content to determine if it contains malicious code.'
        },
        'iframe': {
            'title': 'Embedded iFrame Element',
            'description': 'iFrames embedded in SVG files can load external HTML content. This is often used to redirect users to phishing sites or malicious content.',
            'impact': 'HIGH',
            'mitre_technique': 'T1189 (Drive-by Compromise)',
            'remediation': 'Remove iframe elements from SVG files.'
        },
        'http|https': {
            'title': 'External Resource Reference',
            'description': 'URLs in SVG files can load external resources that may contain malicious content or be used for data exfiltration.',
            'impact': 'MEDIUM',
            'mitre_technique': 'T1071 (Application Layer Protocol)',
            'remediation': 'Remove or sanitize external URLs in SVG files.'
        },
        'fetch|XMLHttpRequest': {
            'title': 'Network Request Method',
            'description': 'JavaScript methods for making HTTP requests. Can be used to exfiltrate data or download additional malicious content.',
            'impact': 'HIGH',
            'mitre_technique': 'T1071.001 (Application Layer Protocol: Web Protocols)',
            'remediation': 'Remove network request code from SVG files.'
        }
    }
    
    # Find the matching explanation based on the pattern
    matched_explanation = None
    for key, explanation in explanations.items():
        if re.search(key, pattern, re.IGNORECASE) or re.search(key, match, re.IGNORECASE):
            matched_explanation = explanation
            break
    
    # If we found a matching explanation
    if matched_explanation:
        return f"""
{indicator_type} Indicator - {matched_explanation['title']} [Impact: {matched_explanation['impact']}]
  Pattern: {pattern}
  Match: {match}
  Description: {matched_explanation['description']}
  MITRE ATT&CK: {matched_explanation['mitre_technique']}
  Recommendation: {matched_explanation['remediation']}"""
    
    # Default explanation if pattern isn't specifically documented
    return f"""
{indicator_type} Indicator - Unknown Pattern [Impact: MEDIUM]
  Pattern: {pattern}
  Match: {match}
  Description: This pattern may indicate malicious behavior or script injection.
    Further investigation is recommended.
  Recommendation: Review the context of this match to determine if it's legitimate."""

def interpret_entropy(entropy):
    """
    Interpret the entropy value in terms of potential file obfuscation.
    
    Args:
        entropy (float): Entropy value of file
        
    Returns:
        str: Interpretation of the entropy value
    """
    if entropy is None or not isinstance(entropy, (int, float)):
        return "Unknown"
    
    if entropy < 1:
        return "Very Low - Likely not obfuscated/encoded"
    elif entropy < 3:
        return "Low - Probably normal text or code"
    elif entropy < 5:
        return "Medium - Normal for mixed content files"
    elif entropy < 7:
        return "High - May contain compressed or encoded data"
    else:
        return "Very High - Likely contains encrypted, compressed, or random data"

def calculate_risk_level(findings):
    """
    Calculate overall risk level based on findings.
    
    Args:
        findings (list): List of findings tuples
        
    Returns:
        str: Risk level assessment
    """
    impact_scores = {
        'CRITICAL': 100,
        'HIGH': 75,
        'MEDIUM': 50,
        'LOW': 25
    }
    
    if not findings:
        return "LOW"
    
    # Initialize risk score
    total_score = 0
    count = 0
    
    # Analyze each finding and get its explanation
    for finding in findings:
        indicator_type, pattern, matches = finding
        for match in matches:
            explanation = explain_threat_indicator(indicator_type, pattern, match)
            
            # Extract impact from explanation
            impact_match = re.search(r'Impact: (\w+)', explanation)
            if impact_match:
                impact = impact_match.group(1)
                total_score += impact_scores.get(impact, 25)  # Default to LOW if not found
                count += 1
    
    # Calculate average score
    if count > 0:
        avg_score = total_score / count
    else:
        avg_score = 0
    
    # Convert score to risk level
    if avg_score >= 80:
        return "CRITICAL"
    elif avg_score >= 60:
        return "HIGH"
    elif avg_score >= 40:
        return "MEDIUM"
    else:
        return "LOW"

# Entrypoint
if __name__ == "__main__":
    main()
